"""Tezin tüm Markdown bölümlerini tek profesyonel Word (.docx) belgesine birleştirir.

Çıktı: Tez_GAN_Urban_Design.docx
Stil: Fırat Üniversitesi Fen Bilimleri Enstitüsü tipik şablonuna yakın
  - Times New Roman 12pt gövde, 1.5 satır aralığı
  - Başlıklar bold, numaralı (1. GİRİŞ → 6. SONUÇLAR)
  - Tablo / Şekil otomatik referans
  - PNG görseller content stream'ine gömülür
  - Türkçe karakter desteği tam (UTF-8)

Kullanım:
  python3 scripts/build_thesis_docx.py
"""

from __future__ import annotations

import hashlib
import os
import re
import shutil
import subprocess
import tempfile
import zipfile
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement, parse_xml


REPO = Path(__file__).resolve().parent.parent
DOCS = REPO / "docs"
FIGS = REPO / "figures"
EQ_CACHE = REPO / "figures" / "_omml"   # OMML XML önbelleği
EQ_CACHE.mkdir(parents=True, exist_ok=True)
OUT  = REPO / "Tez_GAN_Urban_Design.docx"

# Pandoc bul (sistem PATH veya ~/.local/bin)
PANDOC = shutil.which("pandoc") or os.path.expanduser("~/.local/bin/pandoc")
if not os.path.isfile(PANDOC):
    raise SystemExit(
        f"pandoc bulunamadı: {PANDOC}\n"
        "Kurulum: https://github.com/jgm/pandoc/releases -> arm64-macOS.zip indirin, "
        "binary'yi ~/.local/bin/pandoc altına kopyalayın."
    )

# OMML namespace (Word matematik)
OMML_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"


# ---------------------------------------------------------------------------
# LaTeX → OMML (Word native matematik) dönüşümü
# ---------------------------------------------------------------------------
# Pandoc, LaTeX matematiğini OMML'e çevirebilen tek güvenilir araçtır.
# Strateji: tek bir denklem içeren minimal markdown -> docx -> word/document.xml
# içinden <m:oMath ...> XML elementini ayıkla, önbelleğe yaz, sonra
# python-docx ile paragrafa direkt enjekte et.

_TAG_RE = re.compile(r"\\tag\{([^}]+)\}")

_OMATH_PARA_RE = re.compile(
    r"<m:oMathPara[^>]*>(.*?)</m:oMathPara>", re.DOTALL
)
_OMATH_RE = re.compile(
    r"<m:oMath\b[^>]*>.*?</m:oMath>", re.DOTALL
)


def _strip_tag(latex: str) -> tuple[str, str | None]:
    """LaTeX'ten \\tag{N} kısmını ayır."""
    tag = None
    m = _TAG_RE.search(latex)
    if m:
        tag = m.group(1).strip()
        latex = _TAG_RE.sub("", latex).strip()
    return latex, tag


def latex_to_omml(latex: str, display: bool = True) -> tuple[str, str | None]:
    """LaTeX'i pandoc kullanarak OMML XML metnine çevirir.
    Cache'li; aynı denklem için pandoc'u yeniden çağırmaz.
    Döner: (omath_xml_str veya None, tag_no veya None)."""
    formula, tag = _strip_tag(latex)
    key = hashlib.md5(f"{formula}|d{display}".encode("utf-8")).hexdigest()[:14]
    cache_path = EQ_CACHE / f"eq_{key}.xml"

    if cache_path.exists():
        return cache_path.read_text(encoding="utf-8"), tag

    delim = "$$" if display else "$"
    md = f"{delim}{formula}{delim}\n"

    with tempfile.TemporaryDirectory() as td:
        td_path = Path(td)
        in_md = td_path / "in.md"
        out_dx = td_path / "out.docx"
        in_md.write_text(md, encoding="utf-8")
        try:
            subprocess.run(
                [PANDOC, "-f", "markdown", "-t", "docx",
                 str(in_md), "-o", str(out_dx)],
                check=True, capture_output=True, timeout=20,
            )
        except subprocess.CalledProcessError as e:
            print(f"  ! pandoc fail: {formula[:60]!r} -> {e.stderr.decode()[:120]}")
            return None, tag
        except subprocess.TimeoutExpired:
            print(f"  ! pandoc timeout: {formula[:60]!r}")
            return None, tag

        with zipfile.ZipFile(out_dx) as z:
            doc_xml = z.read("word/document.xml").decode("utf-8")

    # Display matematik <m:oMathPara> içine sarmalanır; inline ise sadece <m:oMath>
    m = _OMATH_PARA_RE.search(doc_xml)
    if m:
        omml = "<m:oMathPara xmlns:m=\"%s\">%s</m:oMathPara>" % (
            OMML_NS, m.group(1)
        )
    else:
        m2 = _OMATH_RE.search(doc_xml)
        if not m2:
            return None, tag
        omml = m2.group(0)
        # namespace ekle
        if "xmlns:m=" not in omml:
            omml = omml.replace(
                "<m:oMath", f'<m:oMath xmlns:m="{OMML_NS}"', 1
            )
    cache_path.write_text(omml, encoding="utf-8")
    return omml, tag

BODY_FONT = "Times New Roman"
BODY_SIZE = 12   # pt
HEADING_FONT = "Times New Roman"

# Bölüm sırası (FBE şablon sırası)
SECTIONS = [
    DOCS / "00_On_Sayfalar.md",
    DOCS / "04_Bolum1_Giris.md",
    DOCS / "03_Bolum2_Literatur.md",
    DOCS / "05_Bolum3_Yontem.md",
    DOCS / "06_Bolum4_Deneysel_Sonuclar.md",
    DOCS / "07_Bolum5_Tartisma.md",
    DOCS / "08_Bolum6_Sonuc.md",
]


# ---------------------------------------------------------------------------
# Yardımcılar
# ---------------------------------------------------------------------------
def set_cell_border(cell, color="888888", size="6"):
    """Tablo hücre kenarlık çiz (gri ince)."""
    tc_pr = cell._tc.get_or_add_tcPr()
    # Mevcut tcBorders'ı bul veya oluştur (tek seferlik)
    tcBorders = tc_pr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tc_pr.append(tcBorders)
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), size)
        b.set(qn("w:color"), color)
        tcBorders.append(b)


def set_run_font(run, size=BODY_SIZE, bold=False, italic=False,
                 font=BODY_FONT, color=None):
    run.font.name = font
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    # Türkçe için font ailesi de set edilmeli (Word'de bazı sürümler ek attribute ister)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:cs"), font)


# Inline markdown parsing — **bold**, *italic*, `code`, [link](url), $math$
INLINE_RE = re.compile(
    r"(\*\*[^*]+?\*\*|\*[^*\n]+?\*|`[^`]+?`|\[[^\]]+?\]\([^)]+?\)|\$[^$]+?\$)"
)


def add_inline_runs(paragraph, text, base_size=BODY_SIZE, base_italic=False):
    """Bir markdown satırını parse edip inline run'ları paragrafa ekler."""
    parts = INLINE_RE.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, size=base_size, bold=True, italic=base_italic)
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, size=base_size, italic=True)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, size=base_size - 1, italic=False,
                         font="Courier New", color=RGBColor(0x55, 0x55, 0x55))
        elif part.startswith("[") and "](" in part and part.endswith(")"):
            label, url = re.match(r"\[([^\]]+)\]\(([^)]+)\)", part).groups()
            run = paragraph.add_run(label)
            set_run_font(run, size=base_size, color=RGBColor(0x10, 0x4E, 0x8B))
            run.font.underline = True
        elif part.startswith("$") and part.endswith("$"):
            # Inline matematik: native Word OMML olarak eklenir
            latex = part[1:-1]
            omml, _ = latex_to_omml(latex, display=False)
            if omml:
                # Inline OMath, paragrafın doğrudan çocuğu olarak insert
                # (run sıralamasını koruyarak)
                try:
                    el = parse_xml(omml)
                    paragraph._p.append(el)
                    continue
                except Exception:
                    pass
            # Yedek: ham metin
            run = paragraph.add_run(latex)
            set_run_font(run, size=base_size, italic=True,
                         font="Cambria Math")
        else:
            run = paragraph.add_run(part)
            set_run_font(run, size=base_size, italic=base_italic)


# ---------------------------------------------------------------------------
# Document setup
# ---------------------------------------------------------------------------
def setup_styles(doc: Document):
    # Varsayılan stil
    style = doc.styles["Normal"]
    style.font.name = BODY_FONT
    style.font.size = Pt(BODY_SIZE)
    pf = style.paragraph_format
    pf.line_spacing = 1.5
    pf.space_after = Pt(6)

    # Türkçe ek font attr
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), BODY_FONT)
    rFonts.set(qn("w:hAnsi"), BODY_FONT)


def add_page_break(doc: Document):
    doc.add_paragraph().add_run().add_break()
    from docx.enum.text import WD_BREAK
    last = doc.paragraphs[-1]
    last.add_run().add_break(WD_BREAK.PAGE)


def add_heading(doc: Document, text: str, level: int):
    """Numaralı/numarasız başlık ekler."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18 - 2 * level)
    p.paragraph_format.space_after = Pt(12 - 2 * level)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    sizes = {1: 18, 2: 14, 3: 12.5, 4: 12}
    set_run_font(run, size=sizes.get(level, 12), bold=True)


# ---------------------------------------------------------------------------
# Markdown parser
# ---------------------------------------------------------------------------
FIGURE_REF_RE = re.compile(r"`figures/([^`]+\.png)`")


def parse_markdown_to_docx(md_text: str, doc: Document):
    lines = md_text.split("\n")
    i = 0
    n = len(lines)

    while i < n:
        line = lines[i]
        stripped = line.rstrip()

        # 1) Başlıklar (# .. ######)
        m_h = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if m_h:
            level = len(m_h.group(1))
            text = m_h.group(2).strip()
            add_heading(doc, text, level)
            i += 1
            continue

        # 2) Yatay çizgi
        if stripped == "---":
            p = doc.add_paragraph()
            p_borders = OxmlElement("w:pBdr")
            bot = OxmlElement("w:bottom")
            bot.set(qn("w:val"), "single")
            bot.set(qn("w:sz"), "6")
            bot.set(qn("w:color"), "AAAAAA")
            p_borders.append(bot)
            p.paragraph_format.element.get_or_add_pPr().append(p_borders)
            i += 1
            continue

        # 2.5) Blok denklem $$ ... $$  → Word native OMML (oMathPara)
        if stripped.startswith("$$"):
            if stripped.endswith("$$") and len(stripped) > 4:
                latex = stripped[2:-2].strip()
                i += 1
            else:
                inner = [stripped[2:]]
                i += 1
                while i < n:
                    cur = lines[i].rstrip()
                    if cur.endswith("$$"):
                        inner.append(cur[:-2])
                        i += 1
                        break
                    inner.append(cur)
                    i += 1
                latex = "\n".join(inner).strip()

            omml, tag = latex_to_omml(latex, display=True)
            inserted = False
            if omml:
                try:
                    # Tab-stop: ortada formül, sağa yaslı denklem numarası
                    # OMath kendini ortalar; tag'i ayrı bir paragraf ile değil
                    # tek paragrafta sağa yaslı tab ile ekliyoruz.
                    p = doc.add_paragraph()
                    p.paragraph_format.space_before = Pt(6)
                    p.paragraph_format.space_after = Pt(6)
                    # OMathPara/oMath element'i paragraf çocuğu olarak ekle
                    el = parse_xml(omml)
                    p._p.append(el)
                    if tag:
                        # Numarayı paragraf sonuna sağa yaslı koy (oMath sonrası
                        # ekstra run + sağ tab stop)
                        from docx.enum.text import WD_TAB_ALIGNMENT
                        from docx.shared import Cm as _Cm
                        # Tab stop ekleme
                        tab_stops = p.paragraph_format.tab_stops
                        try:
                            tab_stops.add_tab_stop(_Cm(16.0),
                                                   alignment=WD_TAB_ALIGNMENT.RIGHT)
                        except Exception:
                            pass
                        r = p.add_run(f"\t({tag})")
                        set_run_font(r, size=BODY_SIZE)
                    inserted = True
                except Exception as e:
                    print(f"  ! OMML parse fail: {e}")
            if not inserted:
                # Yedek: salt metin (asla görsel kullanma — tez kuralı gereği)
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(latex + (f"   ({tag})" if tag else ""))
                set_run_font(run, size=BODY_SIZE, italic=True,
                             font="Cambria Math")
            continue

        # 3) Kod bloğu ```...```
        if stripped.startswith("```"):
            i += 1
            code_lines = []
            while i < n and not lines[i].rstrip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1  # closing ```
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.5)
            p.paragraph_format.line_spacing = 1.15
            p.paragraph_format.space_after = Pt(6)
            # Gri arka plan (XML)
            shading = OxmlElement("w:shd")
            shading.set(qn("w:val"), "clear")
            shading.set(qn("w:color"), "auto")
            shading.set(qn("w:fill"), "F2F2F2")
            p.paragraph_format.element.get_or_add_pPr().append(shading)
            run = p.add_run("\n".join(code_lines))
            set_run_font(run, size=10, font="Courier New",
                         color=RGBColor(0x33, 0x33, 0x33))
            continue

        # 4) Blockquote >
        if stripped.startswith(">"):
            block = []
            while i < n and lines[i].rstrip().startswith(">"):
                block.append(lines[i].rstrip().lstrip(">").lstrip())
                i += 1
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.8)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            # Sol kenar çubuğu
            pbd = OxmlElement("w:pBdr")
            left = OxmlElement("w:left")
            left.set(qn("w:val"), "single")
            left.set(qn("w:sz"), "12")
            left.set(qn("w:color"), "888888")
            left.set(qn("w:space"), "4")
            pbd.append(left)
            p.paragraph_format.element.get_or_add_pPr().append(pbd)
            add_inline_runs(p, " ".join(block), base_italic=True)
            continue

        # 5) Tablo (| col1 | col2 |)
        if stripped.startswith("|") and stripped.endswith("|"):
            table_lines = []
            while i < n and lines[i].rstrip().startswith("|"):
                table_lines.append(lines[i].rstrip())
                i += 1
            render_table(doc, table_lines)
            continue

        # 6) Liste (markdown: - veya 1.)
        m_ul = re.match(r"^(\s*)[-•]\s+(.*)$", stripped)
        m_ol = re.match(r"^(\s*)(\d+)\.\s+(.*)$", stripped)
        if m_ul or m_ol:
            list_block = []
            while i < n:
                cur = lines[i].rstrip()
                if re.match(r"^(\s*)[-•]\s+", cur) or re.match(r"^(\s*)\d+\.\s+", cur):
                    list_block.append(cur)
                    i += 1
                elif cur.strip() == "" and (i + 1 < n) and (
                    re.match(r"^(\s*)[-•]\s+", lines[i+1].rstrip()) or
                    re.match(r"^(\s*)\d+\.\s+", lines[i+1].rstrip())
                ):
                    i += 1   # listeyi sürdür
                else:
                    break
            for li in list_block:
                m1 = re.match(r"^(\s*)([-•])\s+(.*)$", li)
                m2 = re.match(r"^(\s*)(\d+)\.\s+(.*)$", li)
                indent_chars = (m1 or m2).group(1)
                indent = Cm(0.6 + 0.6 * (len(indent_chars) // 2))
                content = (m1 or m2).group(3)
                style_name = "List Bullet" if m1 else "List Number"
                try:
                    p = doc.add_paragraph(style=style_name)
                except KeyError:
                    p = doc.add_paragraph()
                p.paragraph_format.left_indent = indent
                p.paragraph_format.space_after = Pt(3)
                if not m1 and not m2:
                    continue
                if not p.text and m1:
                    # marker run zaten ekleniyorsa atla
                    pass
                # bullet stilini varsayılan koru
                add_inline_runs(p, content)
            continue

        # 7) Şekil referansı satırı (`figures/...png`)
        m_fig = FIGURE_REF_RE.search(stripped)
        if m_fig and ("Şekil" in stripped or "Görsel" in stripped or stripped.startswith("![")):
            fname = m_fig.group(1)
            insert_figure(doc, fname, caption=stripped)
            i += 1
            continue

        # 8) Boş satır
        if stripped == "":
            i += 1
            continue

        # 9) Çoklu satır paragrafı — boş satıra kadar topla
        para_lines = [stripped]
        i += 1
        while i < n:
            nxt = lines[i].rstrip()
            if nxt == "":
                break
            if re.match(r"^#{1,6}\s+", nxt): break
            if nxt.startswith("|"): break
            if nxt.startswith(">"): break
            if nxt.startswith("```"): break
            if nxt.startswith("$$"): break
            if re.match(r"^(\s*)[-•]\s+", nxt): break
            if re.match(r"^(\s*)\d+\.\s+", nxt): break
            if nxt == "---": break
            para_lines.append(nxt)
            i += 1
        text = " ".join(para_lines).strip()
        # Şekil referansı satırı paragraf içinde gizliyse, görseli ekle
        if FIGURE_REF_RE.search(text) and "Şekil" in text:
            m = FIGURE_REF_RE.search(text)
            insert_figure(doc, m.group(1), caption=text)
        else:
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Cm(0.6)
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            add_inline_runs(p, text)


def render_table(doc: Document, table_lines):
    """Markdown tablosunu docx tablosu olarak render eder."""
    # Header ayraç satırını (|---|---|) at
    rows = []
    for line in table_lines:
        if re.match(r"^\|\s*:?-+:?\s*(\|\s*:?-+:?\s*)+\|", line):
            continue
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        rows.append(cells)
    if not rows:
        return
    ncol = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncol)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for r_idx, cells in enumerate(rows):
        for c_idx in range(ncol):
            cell = table.cell(r_idx, c_idx)
            txt = cells[c_idx] if c_idx < len(cells) else ""
            # Var olan boş paragrafı kaldır
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.15
            add_inline_runs(p, txt, base_size=10)
            # Header arka plan
            if r_idx == 0:
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), "DDE6F0")
                cell._tc.get_or_add_tcPr().append(shd)
                for run in p.runs:
                    run.bold = True
            set_cell_border(cell)
    # Tablo sonrası boş satır
    doc.add_paragraph()


def insert_figure(doc: Document, filename: str, caption: str = ""):
    img_path = FIGS / filename
    if not img_path.exists():
        # caption metnini düşür ama görsel olmadığını belirt
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f"[Görsel bulunamadı: {filename}]")
        set_run_font(r, italic=True, color=RGBColor(0xAA, 0x44, 0x44))
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run()
    try:
        run.add_picture(str(img_path), width=Cm(15.0))
    except Exception as e:
        run.text = f"[Görsel yüklenemedi: {e}]"
    # Caption — Şekil numarası varsa ön plana çıkar
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_after = Pt(12)
        # caption'dan ham filename ref'ini temizle
        clean = re.sub(r"\s*\(görsel:\s*`figures/[^`]+`\)\.?", "", caption)
        clean = re.sub(r"\*\*", "", clean).strip()
        r = cap.add_run(clean)
        set_run_font(r, size=10, italic=True, color=RGBColor(0x44, 0x44, 0x44))


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------
def main():
    doc = Document()
    setup_styles(doc)

    # Sayfa kenar boşlukları (FBE: sol 3 cm, sağ 2.5 cm, üst 3 cm, alt 2.5 cm)
    for section in doc.sections:
        section.top_margin = Cm(3.0)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.5)

    # Kapak sayfası (placeholder)
    cover = doc.add_paragraph()
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cover.add_run("\n\n\n\nT.C.\nFIRAT ÜNİVERSİTESİ\nFEN BİLİMLERİ ENSTİTÜSÜ")
    set_run_font(r, size=14, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("\n\n\nGAN MODELLERİ KULLANARAK\nKENTSEL PEYZAJ TASARIM PLANLARININ\nOTOMATİK OLUŞTURULMASI VE RENKLENDİRİLMESİ")
    set_run_font(r, size=16, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("\n\n\nYÜKSEK LİSANS TEZİ")
    set_run_font(r, size=13, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("\n\n[Öğrenci Adı Soyadı]\n[Anabilim Dalı]")
    set_run_font(r, size=12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("\n\n\nDanışman: [Unvan + Ad Soyad]\n\nELAZIĞ — 2026")
    set_run_font(r, size=12, bold=True)

    # Sayfa atlatma
    doc.add_page_break()

    # Her bölümü işle
    for src in SECTIONS:
        if not src.exists():
            print(f"  UYARI: {src} bulunamadı, atlandı.")
            continue
        print(f"  -> {src.name}")
        text = src.read_text(encoding="utf-8")
        # Front-matter (eğer varsa) atla
        if text.startswith("---"):
            parts = text.split("---", 2)
            if len(parts) >= 3:
                text = parts[2]
        parse_markdown_to_docx(text, doc)
        # Her bölümden sonra sayfa atla (00 hariç)
        if src.name != "00_On_Sayfalar.md":
            doc.add_page_break()

    # Görseller bölümü — figures/ tüm PNG'leri özet ek olarak ekle
    doc.add_page_break()
    add_heading(doc, "EK A — Tezde Üretilen Görseller", level=1)
    intro = doc.add_paragraph()
    intro.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_inline_runs(intro,
        "Bu ekte tezde kullanılan statik görseller (mimari diyagramları ve niceliksel grafikler) "
        "yüksek çözünürlüklü olarak yeniden sunulmaktadır. Tüm görseller `figures/` dizininde "
        "PNG (300 DPI) olarak commit edilmiştir; yeniden üretim için "
        "`python3 scripts/figures/make_static_figures.py` çalıştırılır.")

    fig_order = [
        ("fig_3_1_pipeline.png",       "Şekil 3.1. Önerilen iki aşamalı GAN pipeline'ı."),
        ("fig_3_2_unet.png",            "Şekil 3.2. U-Net jeneratör mimarisi (Pix2Pix / Geliştirilmiş Pix2Pix)."),
        ("fig_3_3_patchgan.png",        "Şekil 3.3. 70x70 PatchGAN ayırıcı mimarisi."),
        ("fig_3_4_spade.png",           "Şekil 3.4. SPADE (uzamsal-uyarlamalı denormalizasyon) blok diyagramı."),
        ("fig_4_3_metrics_bars.png",    "Şekil 4.3. Üç GAN mimarisinin beş metrik üzerinde performans karşılaştırması."),
        ("fig_4_4_radar.png",           "Şekil 4.4. Modellerin normalize metrik profili (radar)."),
        ("fig_4_5_improvements.png",    "Şekil 4.5. Geliştirilmiş Pix2Pix'in baseline'a göre metrik bazlı iyileşmesi."),
        ("fig_4_6_tradeoff.png",        "Şekil 4.6. Algı-bozulma ödünleşimi ve Geliştirilmiş Pix2Pix'in aşması."),
    ]
    for fname, cap in fig_order:
        insert_figure(doc, fname, caption=cap)

    print(f"\n>> Kaydediliyor: {OUT}")
    doc.save(str(OUT))
    print(f">> Hazır: {OUT}  ({OUT.stat().st_size/1024:.1f} KB)")


if __name__ == "__main__":
    main()
